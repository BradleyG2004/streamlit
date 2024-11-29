import streamlit as st
from transformers import pipeline
import whisper
from pyannote.audio import Pipeline
from moviepy.editor import VideoFileClip
from pydub import AudioSegment
import librosa
import noisereduce as nr
import soundfile as sf
import spacy
from unidecode import unidecode
from docx import Document
import os
from langchain.prompts import ChatPromptTemplate
from langchain_community.llms.ollama import Ollama

# Charger les modèles
model_id = "llama3.1"
model = Ollama(model=model_id)
model1 = whisper.load_model("medium")
nlp = spacy.load("xx_ent_wiki_sm")
# Fonction pour extraire le nom du locuteur
def extract_speaker_name(text):
    doc = nlp(text)
    names = [ent.text for ent in doc.ents if ent.label_ == "PER"]
    if names:
        return names[0]
    return None

# Fonction pour enlever les accents
def remove_accents(text):
    return unidecode(text)

# Fonction de transcription
def transcribe(audio_path, lang):
    if lang != 'Auto':
        trans = model1.transcribe(audio_path, fp16=False, language=lang)
    else:
        trans = model1.transcribe(audio_path, fp16=False)
    return trans['segments']

# Fonction d'exportation de la transcription
def export_transcription_only(transcription_output, filename="transcription.docx"):
    doc = Document()
    doc.add_heading("Transcription", level=1)
    for line in transcription_output.split("\n"):
        doc.add_paragraph(line)
    doc.save(filename)
    return filename
def export_summary_only(summary, filename="résumé.docx"):
    doc = Document()
    doc.add_heading("Résumé", level=1)
    doc.add_paragraph(summary)
    doc.save(filename)
    return filename

# Fonction d'exportation de la transcription et du résumé
def export_transcription_and_summary(transcription, summary, filename="transcription_et_resume.docx"):
    doc = Document()
    doc.add_heading("Transcription et Résumé", level=1)

    # Ajouter la transcription
    doc.add_heading("Transcription", level=2)
    for line in transcription.split('\n'):
        doc.add_paragraph(line)

    # Ajouter le résumé
    doc.add_heading("Résumé", level=2)
    doc.add_paragraph(summary)

    # Sauvegarder le fichier
    doc.save(filename)
    return filename

# Fonction de résumé
def summarize_text(input_text, language):
    prompt_template = """
    You are an advanced summarization model. Summarize the following text in {language}:
    
    {text}
    
    Provide a concise and informative summary. Make sure to check your result to see if it is grammatically correct (No need to communicate it in the final result).
    """
    prompt = ChatPromptTemplate.from_template(prompt_template)
    input_prompt = prompt.format(text=input_text, language=language)
    result = model(input_prompt)
    return result.strip()

# Interface Streamlit
st.title("Outil de synthèse de meetings en ligne")

# Langues supportées
supported_languages = ['fr', 'en', 'es', 'de', 'it', 'pt', 'ru', 'zh', 'ja', 'ko']
lang = st.selectbox("Choisissez la langue de transcription", supported_languages + ['Auto'])

# Initialiser la session
if "summary" not in st.session_state:
    st.session_state["summary"] = ""

# Téléchargement du fichier
uploaded_file = st.file_uploader("Veuillez uploader un fichier Audio/Video (.mp3, .wav, .mp4, .mpeg4)", type=["mp3", "wav", "mp4", "mpeg4"])

if uploaded_file:
    with st.spinner('Transcription en cours...'):
        # Préparation du fichier audio
        if uploaded_file.type in ["audio/mpeg", "audio/wav"]:
            audio = AudioSegment.from_file(uploaded_file)
            wav_path = "audio_temp.wav"
            audio.export(wav_path, format="wav")
        elif uploaded_file.type == "video/mp4":
            video = VideoFileClip(uploaded_file.name)
            wav_path = "audio_extracted.wav"
            video.audio.write_audiofile(wav_path)
        else:
            st.error("Fichier non valide.")
            st.stop()

        # Réduction de bruit
        signal, sr = librosa.load(wav_path, sr=None)
        reduced_noise = nr.reduce_noise(y=signal, sr=sr)
        sf.write(wav_path, reduced_noise, sr)

        
       # Diarisation des locuteurs
        pipeline = Pipeline.from_pretrained("pyannote/speaker-diarization-3.1", use_auth_token="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")
        diarization = pipeline(wav_path)
        
        merged_segments = []
        current_speaker = None
        current_segment = None

        # Regroupement des segments par locuteur
        for turn, _, speaker in diarization.itertracks(yield_label=True):
            if speaker == current_speaker:
                current_segment = (current_segment[0], turn.end)
            else:
                if current_segment:
                    merged_segments.append((current_segment, current_speaker))
                current_segment = (turn.start, turn.end)
                current_speaker = speaker

        if current_segment:
            merged_segments.append((current_segment, current_speaker))

        final_transcription = []
        for (start, end), speaker in merged_segments:
            segment_path = f"temp_segment_{start:.2f}_{end:.2f}.wav"
            segment_audio = audio[start * 1000:end * 1000]
            segment_audio.export(segment_path, format="wav")
            transcribed_segment = transcribe(segment_path, lang)
            speaker_transcription = ""

            for ts in transcribed_segment:
                speaker_transcription += ts['text'] + " "

            cleaned_transcription = remove_accents(speaker_transcription)
            extracted_name = extract_speaker_name(cleaned_transcription)
            if extracted_name:
                speaker = extracted_name

            duration = end - start
            final_transcription.append({
                "speaker": speaker,
                "start": start,
                "end": end,
                "duration": duration,
                "text": speaker_transcription.strip()
            })
            os.remove(segment_path)

        # Création de la transcription complète
        transcription_output = ""
        transcript = " "
        for i, segment in enumerate(final_transcription):
            transcription_output += (
                f"Début: {segment['start']:.1f}s - Fin: {segment['end']:.1f}s - "
                f"Durée: {segment['duration']:.1f}s - Locuteur: {segment['speaker']}\n"
                f"{segment['text']}\n"
            )
            transcript = " ".join([segment["text"] for segment in final_transcription])
            if i < len(final_transcription) - 1:
                transcription_output += ("---" * 20) + "\n"

        st.text_area("Transcription", transcription_output, height=300, disabled=True)
        st.success('Transcription Complétée')
        transcription_complete = True

        # Exporter uniquement la transcription
        if st.button("Exporter la transcription"):
            file_path = export_transcription_only(transcription_output)
            with open(file_path, "rb") as f:
                st.download_button(
                    label="Télécharger la transcription",
                    data=f,
                    file_name=file_path,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

        # Génération et export du résumé
        if transcription_output:
            summary_lang = st.selectbox("Choisissez la langue du résumé", supported_languages)
            if st.button("Générer le résumé"):
                with st.spinner("Génération du résumé..."):
                    st.session_state["summary"] = summarize_text(transcript, summary_lang)
                    st.subheader("Résumé généré")
                    st.write(st.session_state["summary"])

            #Exporter résumé
            if st.session_state["summary"]:
                if st.button("Exporter résumé"):
                    sum=export_summary_only(st.session_state["summary"])
                    with open(sum, "rb") as f:
                        st.download_button(
                            label="Télécharger le résumé",
                            data=f,
                            file_name=sum,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )


            # Exporter transcription et résumé
            if transcription_output and st.session_state["summary"]:
                if st.button("Exporter transcription et résumé"):
                    combined_filename = export_transcription_and_summary(
                        transcription_output, 
                        st.session_state["summary"]
                    )
                    with open(combined_filename, "rb") as f:
                        st.download_button(
                            label="Télécharger transcription et résumé",
                            data=f,
                            file_name=combined_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )

st.sidebar.header("Audio original")
if uploaded_file:
    st.sidebar.audio(uploaded_file)
